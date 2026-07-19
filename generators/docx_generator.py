"""
DOCX Workbook Generator

Generates professionally formatted DOCX workbooks with optimized
font sizes and paper usage.
"""

from pathlib import Path
from collections import Counter
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from config import (
    OUTPUT_DIR, WORKBOOK_TITLE,
    TITLE_FONT_SIZE, HEADING_FONT_SIZE, QUESTION_FONT_SIZE,
    ANSWER_FONT_SIZE, EXPLANATION_FONT_SIZE, TIP_FONT_SIZE, INFO_FONT_SIZE,
    PAGE_MARGIN_TOP, PAGE_MARGIN_BOTTOM, PAGE_MARGIN_LEFT, PAGE_MARGIN_RIGHT,
    SPACE_BEFORE_QUESTION, SPACE_AFTER_QUESTION,
    SPACE_BEFORE_ANSWER, SPACE_AFTER_ANSWER,
    SPACE_BEFORE_EXPLANATION, SPACE_AFTER_EXPLANATION,
    SPACE_BEFORE_TIP, SPACE_AFTER_TIP,
    SPACE_BEFORE_SEPARATOR, SPACE_AFTER_SEPARATOR,
    QUESTIONS_PER_PAGE
)
from models.question import Question


class DOCXGenerator:
    """Generates a professionally formatted DOCX workbook."""

    def __init__(self, questions: list[Question]) -> None:
        self.questions = questions
        self.output_file = OUTPUT_DIR / "SQL_Workbook.docx"

    def set_cell_border(self, cell, **kwargs):
        """Set cell border for table formatting."""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        
        for border_name in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '999999')
            tcPr.append(border)

    def generate(self) -> Path:
        """Generate the DOCX workbook with optimized formatting."""
        
        doc = Document()
        
        # Set page margins
        section = doc.sections[0]
        section.left_margin = Inches(PAGE_MARGIN_LEFT)
        section.right_margin = Inches(PAGE_MARGIN_RIGHT)
        section.top_margin = Inches(PAGE_MARGIN_TOP)
        section.bottom_margin = Inches(PAGE_MARGIN_BOTTOM)
        
        # ============================================================
        # Title Page
        # ============================================================
        
        title = doc.add_heading(WORKBOOK_TITLE, level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.runs[0].font.size = Pt(TITLE_FONT_SIZE)
        title.runs[0].font.bold = True
        title.runs[0].font.color.rgb = RGBColor(102, 126, 234)
        
        subtitle = doc.add_paragraph(f"{len(self.questions)} SQL Practice Questions")
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.runs[0].font.size = Pt(20)
        subtitle.runs[0].font.italic = True
        
        doc.add_paragraph()
        
        # Stats
        stats = doc.add_paragraph()
        stats.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        diff_counts = Counter([q.difficulty for q in self.questions])
        stats_text = " | ".join([f"{diff}: {count}" for diff, count in diff_counts.items()])
        stats_run = stats.add_run(stats_text)
        stats_run.font.size = Pt(14)
        
        # Note about answers
        note = doc.add_paragraph()
        note.alignment = WD_ALIGN_PARAGRAPH.CENTER
        note_run = note.add_run("📚 Answers with explanations available in Answers.sql file")
        note_run.font.size = Pt(12)
        note_run.font.italic = True
        
        doc.add_page_break()
        
        # ============================================================
        # Table of Contents
        # ============================================================
        
        toc_title = doc.add_heading("Table of Contents", level=1)
        toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        toc_title.runs[0].font.size = Pt(24)
        
        toc_table = doc.add_table(rows=1, cols=3)
        toc_table.style = 'Table Grid'
        toc_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        toc_table.autofit = False
        toc_table.columns[0].width = Inches(0.8)
        toc_table.columns[1].width = Inches(3.8)
        toc_table.columns[2].width = Inches(1.4)
        
        header_cells = toc_table.rows[0].cells
        header_cells[0].text = "Q#"
        header_cells[1].text = "Topic"
        header_cells[2].text = "Difficulty"
        
        for cell in header_cells:
            cell.paragraphs[0].runs[0].font.size = Pt(14)
            cell.paragraphs[0].runs[0].font.bold = True
        
        max_toc = min(100, len(self.questions))
        for i in range(max_toc):
            q = self.questions[i]
            row_cells = toc_table.add_row().cells
            row_cells[0].text = str(q.id)
            row_cells[1].text = q.topic[:40] + ('...' if len(q.topic) > 40 else '')
            row_cells[2].text = q.difficulty
            
            for cell in row_cells:
                cell.paragraphs[0].runs[0].font.size = Pt(11)
        
        if len(self.questions) > 100:
            row_cells = toc_table.add_row().cells
            row_cells[0].text = "..."
            row_cells[1].text = f"and {len(self.questions) - 100} more questions"
            row_cells[2].text = ""
            for cell in row_cells:
                cell.paragraphs[0].runs[0].font.size = Pt(11)
        
        # Note about answers
        answer_note = doc.add_paragraph()
        answer_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ans_note_run = answer_note.add_run("📚 For detailed answers with explanations, see Answers.sql file")
        ans_note_run.font.size = Pt(12)
        ans_note_run.font.italic = True
        
        doc.add_page_break()
        
        # ============================================================
        # Questions - Optimized Layout
        # ============================================================
        
        for idx, question in enumerate(self.questions, 1):
            # Question number
            q_heading = doc.add_heading(f"Question {question.id}", level=2)
            q_heading.runs[0].font.size = Pt(HEADING_FONT_SIZE)
            q_heading.runs[0].font.bold = True
            q_heading.runs[0].font.color.rgb = RGBColor(102, 126, 234)
            q_heading.paragraph_format.space_after = Pt(SPACE_AFTER_QUESTION)
            q_heading.paragraph_format.space_before = Pt(SPACE_BEFORE_QUESTION)
            
            # Topic and Difficulty
            info_p = doc.add_paragraph()
            info_p.paragraph_format.space_after = Pt(2)
            
            info_p.add_run("Topic: ").bold = True
            info_p.add_run(question.topic)
            info_p.add_run("    ")
            
            info_p.add_run("Difficulty: ").bold = True
            stars = {
                'Beginner': '⭐',
                'Intermediate': '⭐⭐',
                'Advanced': '⭐⭐⭐',
                'Expert': '⭐⭐⭐⭐'
            }
            info_p.add_run(f"{question.difficulty} {stars.get(question.difficulty, '')}")
            
            for run in info_p.runs:
                run.font.size = Pt(INFO_FONT_SIZE)
            
            # Question
            q_p = doc.add_paragraph()
            q_p.paragraph_format.space_after = Pt(4)
            q_p.paragraph_format.space_before = Pt(4)
            q_run = q_p.add_run(f"Question: {question.question}")
            q_run.font.size = Pt(QUESTION_FONT_SIZE)
            q_run.font.bold = True
            
            # Note about answers
            answer_ref = doc.add_paragraph()
            answer_ref.paragraph_format.space_after = Pt(4)
            answer_ref.paragraph_format.space_before = Pt(2)
            ref_run = answer_ref.add_run("📚 See Answers.sql for detailed solution with explanation")
            ref_run.font.size = Pt(10)
            ref_run.font.italic = True
            ref_run.font.color.rgb = RGBColor(100, 100, 100)
            
            # Separator
            sep_p = doc.add_paragraph()
            sep_p.paragraph_format.space_after = Pt(SPACE_AFTER_SEPARATOR)
            sep_p.paragraph_format.space_before = Pt(SPACE_BEFORE_SEPARATOR)
            sep_run = sep_p.add_run("─" * 80)
            sep_run.font.size = Pt(8)
            sep_run.font.color.rgb = RGBColor(200, 200, 200)
            
            # Page break every N questions
            if idx % QUESTIONS_PER_PAGE == 0 and idx != len(self.questions):
                doc.add_page_break()
        
        # Footer note
        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer.add_run("📚 Complete answers with explanations are available in Answers.sql")
        footer_run.font.size = Pt(12)
        footer_run.font.italic = True
        
        doc.save(self.output_file)
        return self.output_file
