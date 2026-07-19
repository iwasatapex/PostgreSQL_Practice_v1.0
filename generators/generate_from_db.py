#!/usr/bin/env python3
"""
Generate SQL Workbook from Database Data
Creates realistic questions based on actual data in your PostgreSQL database
"""

import os
import json
import random
import psycopg2
from datetime import datetime
from pathlib import Path

# ============================================================
# DATABASE CONNECTION
# ============================================================

DB_SETTINGS = {
    "dbname": "postgres",
    "user": "postgres",
    "password": "4m4teur",
    "host": "localhost",
    "port": "5432"
}

class DatabaseConnector:
    def __init__(self):
        self.conn = None
        self.cursor = None
    
    def connect(self):
        try:
            self.conn = psycopg2.connect(**DB_SETTINGS)
            self.cursor = self.conn.cursor()
            return True
        except Exception as e:
            print(f"❌ Connection failed: {e}")
            return False
    
    def disconnect(self):
        if self.cursor:
            self.cursor.close()
        if self.conn:
            self.conn.close()
    
    def get_tables(self):
        self.cursor.execute("""
            SELECT table_name 
            FROM information_schema.tables 
            WHERE table_schema='public'
            ORDER BY table_name;
        """)
        return [row[0] for row in self.cursor.fetchall()]
    
    def get_table_info(self, table_name):
        self.cursor.execute("""
            SELECT 
                column_name,
                data_type,
                is_nullable = 'YES' as nullable,
                (SELECT COUNT(*) FROM information_schema.key_column_usage 
                 WHERE table_name = %s AND column_name = c.column_name 
                 AND constraint_name LIKE '%pkey%') as primary_key,
                (SELECT COUNT(*) FROM information_schema.key_column_usage 
                 WHERE table_name = %s AND column_name = c.column_name 
                 AND constraint_name LIKE '%fkey%') as foreign_key
            FROM information_schema.columns c
            WHERE table_schema='public' AND table_name = %s
            ORDER BY ordinal_position;
        """, (table_name, table_name, table_name))
        
        columns = []
        for row in self.cursor.fetchall():
            columns.append({
                'name': row[0],
                'data_type': row[1],
                'nullable': row[2],
                'primary_key': row[3] > 0,
                'foreign_key': row[4] > 0
            })
        return columns
    
    def get_row_count(self, table_name):
        self.cursor.execute(f"SELECT COUNT(*) FROM {table_name};")
        return self.cursor.fetchone()[0]
    
    def get_sample_data(self, table_name, limit=5):
        self.cursor.execute(f"SELECT * FROM {table_name} LIMIT {limit};")
        columns = [desc[0] for desc in self.cursor.description]
        rows = self.cursor.fetchall()
        return [dict(zip(columns, row)) for row in rows]
    
    def get_column_values(self, table_name, column_name, limit=20):
        self.cursor.execute(f"""
            SELECT DISTINCT {column_name} 
            FROM {table_name} 
            WHERE {column_name} IS NOT NULL 
            LIMIT {limit};
        """)
        return [row[0] for row in self.cursor.fetchall()]
    
    def get_foreign_keys(self, table_name):
        self.cursor.execute("""
            SELECT
                kcu.column_name,
                ccu.table_name AS foreign_table_name,
                ccu.column_name AS foreign_column_name
            FROM information_schema.table_constraints AS tc
            JOIN information_schema.key_column_usage AS kcu
                ON tc.constraint_name = kcu.constraint_name
                AND tc.table_schema = kcu.table_schema
            JOIN information_schema.constraint_column_usage AS ccu
                ON ccu.constraint_name = tc.constraint_name
                AND ccu.table_schema = tc.table_schema
            WHERE tc.constraint_type = 'FOREIGN KEY'
                AND tc.table_name = %s
                AND tc.table_schema = 'public';
        """, (table_name,))
        return [{'column': row[0], 'ref_table': row[1], 'ref_column': row[2]} for row in self.cursor.fetchall()]

# ============================================================
# QUESTION GENERATOR FROM DATABASE
# ============================================================

class DatabaseQuestionGenerator:
    def __init__(self, db_connector, total_questions=100):
        self.db = db_connector
        self.total = total_questions
        self.tables = {}
        self.questions = []
        
    def load_schema(self):
        """Load complete schema with data samples."""
        table_names = self.db.get_tables()
        
        for table in table_names:
            self.tables[table] = {
                'columns': self.db.get_table_info(table),
                'row_count': self.db.get_row_count(table),
                'sample_data': self.db.get_sample_data(table, 5),
                'foreign_keys': self.db.get_foreign_keys(table)
            }
        
        return self.tables
    
    def generate(self):
        """Generate all questions."""
        self.load_schema()
        
        # Generate various types of questions
        self.generate_select_questions()
        self.generate_where_questions()
        self.generate_order_by_questions()
        self.generate_limit_questions()
        self.generate_distinct_questions()
        self.generate_aggregation_questions()
        self.generate_group_by_questions()
        self.generate_join_questions()
        self.generate_subquery_questions()
        self.generate_complex_questions()
        
        # Shuffle and limit
        random.shuffle(self.questions)
        if len(self.questions) > self.total:
            self.questions = self.questions[:self.total]
        
        # Reassign IDs
        for i, q in enumerate(self.questions, 1):
            q['id'] = i
        
        return self.questions
    
    def add_question(self, topic, difficulty, question, answer, explanation, tip):
        self.questions.append({
            'id': len(self.questions) + 1,
            'topic': topic,
            'difficulty': difficulty,
            'question': question,
            'answer': answer,
            'explanation': explanation,
            'tip': tip
        })
    
    # ============================================================
    # SELECT QUESTIONS
    # ============================================================
    
    def generate_select_questions(self):
        for table, info in self.tables.items():
            if info['row_count'] == 0:
                continue
            
            # SELECT all
            self.add_question(
                topic='SELECT',
                difficulty='Beginner',
                question=f"List all records from the {table} table.",
                answer=f"SELECT * FROM {table};",
                explanation=f"SELECT * returns every column and row from the {table} table.",
                tip="Use specific columns instead of SELECT * in production."
            )
            
            # SELECT specific columns
            cols = [col['name'] for col in info['columns'] if not col['primary_key']][:4]
            if len(cols) >= 2:
                col_list = ', '.join(cols[:3])
                self.add_question(
                    topic='SELECT',
                    difficulty='Beginner',
                    question=f"Display only {', '.join(cols[:3])} columns from the {table} table.",
                    answer=f"SELECT {col_list} FROM {table};",
                    explanation=f"SELECT with specific columns returns only those columns.",
                    tip="Specifying columns improves performance and readability."
                )
    
    # ============================================================
    # WHERE QUESTIONS
    # ============================================================
    
    def generate_where_questions(self):
        for table, info in self.tables.items():
            if info['row_count'] == 0:
                continue
            
            # Text columns with LIKE
            text_cols = [col for col in info['columns'] if 'char' in col['data_type'].lower()]
            for col in text_cols[:2]:
                values = self.db.get_column_values(table, col['name'])
                if values:
                    sample = str(values[0])
                    if len(sample) > 5:
                        sample = sample[:10]
                    self.add_question(
                        topic='WHERE',
                        difficulty='Beginner',
                        question=f"Find all {table} where {col['name']} contains '{sample}'.",
                        answer=f"SELECT * FROM {table} WHERE {col['name']} LIKE '%{sample}%';",
                        explanation=f"LIKE with % performs a partial match on {col['name']}.",
                        tip="Use ILIKE for case-insensitive searches."
                    )
            
            # Numeric columns with comparisons
            num_cols = [col for col in info['columns'] if col['data_type'] in ['integer', 'numeric', 'decimal']]
            for col in num_cols[:2]:
                values = self.db.get_column_values(table, col['name'])
                if values:
                    sample = values[0]
                    self.add_question(
                        topic='WHERE',
                        difficulty='Beginner',
                        question=f"Find all {table} where {col['name']} > {sample}.",
                        answer=f"SELECT * FROM {table} WHERE {col['name']} > {sample};",
                        explanation=f"Filters rows where {col['name']} is greater than {sample}.",
                        tip="Use >, <, >=, <=, =, != for numeric comparisons."
                    )
            
            # NULL checks
            nullable_cols = [col for col in info['columns'] if col['nullable']]
            if nullable_cols:
                col = nullable_cols[0]
                self.add_question(
                    topic='WHERE',
                    difficulty='Beginner',
                    question=f"Find all {table} where {col['name']} is NULL.",
                    answer=f"SELECT * FROM {table} WHERE {col['name']} IS NULL;",
                    explanation=f"IS NULL checks for NULL values in {col['name']}.",
                    tip="Use IS NULL, not = NULL."
                )
    
    # ============================================================
    # ORDER BY QUESTIONS
    # ============================================================
    
    def generate_order_by_questions(self):
        for table, info in self.tables.items():
            if info['row_count'] == 0:
                continue
            
            for col in info['columns']:
                if col['primary_key'] or 'timestamp' in col['data_type'].lower():
                    self.add_question(
                        topic='ORDER BY',
                        difficulty='Beginner',
                        question=f"List all {table} ordered by {col['name']} descending.",
                        answer=f"SELECT * FROM {table} ORDER BY {col['name']} DESC;",
                        explanation=f"ORDER BY {col['name']} DESC sorts from highest to lowest.",
                        tip="DESC shows largest numbers or latest dates first."
                    )
                    break
    
    # ============================================================
    # LIMIT QUESTIONS
    # ============================================================
    
    def generate_limit_questions(self):
        for table, info in self.tables.items():
            if info['row_count'] > 10:
                limit = random.choice([5, 10, 15, 20])
                self.add_question(
                    topic='LIMIT',
                    difficulty='Beginner',
                    question=f"Show only the first {limit} records from the {table} table.",
                    answer=f"SELECT * FROM {table} LIMIT {limit};",
                    explanation=f"LIMIT {limit} restricts output to {limit} rows.",
                    tip="Use LIMIT with OFFSET for pagination."
                )
    
    # ============================================================
    # DISTINCT QUESTIONS
    # ============================================================
    
    def generate_distinct_questions(self):
        for table, info in self.tables.items():
            if info['row_count'] > 10:
                text_cols = [col for col in info['columns'] if 'char' in col['data_type'].lower()]
                if text_cols:
                    col = text_cols[0]
                    self.add_question(
                        topic='DISTINCT',
                        difficulty='Beginner',
                        question=f"Show all unique {col['name']} values from the {table} table.",
                        answer=f"SELECT DISTINCT {col['name']} FROM {table};",
                        explanation=f"DISTINCT removes duplicate values from {col['name']}.",
                        tip="DISTINCT applies to all columns in the SELECT list."
                    )
    
    # ============================================================
    # AGGREGATION QUESTIONS
    # ============================================================
    
    def generate_aggregation_questions(self):
        for table, info in self.tables.items():
            if info['row_count'] == 0:
                continue
            
            num_cols = [col for col in info['columns'] if col['data_type'] in ['integer', 'numeric', 'decimal']]
            for col in num_cols[:1]:
                self.add_question(
                    topic='Aggregation',
                    difficulty='Intermediate',
                    question=f"Calculate the average of {col['name']} from the {table} table.",
                    answer=f"SELECT AVG({col['name']}) AS avg_{col['name']} FROM {table};",
                    explanation=f"AVG({col['name']}) calculates the mean of all {col['name']} values.",
                    tip="Use ROUND(AVG(column), 2) to limit decimal places."
                )
                
                self.add_question(
                    topic='Aggregation',
                    difficulty='Intermediate',
                    question=f"Find the total {col['name']} from the {table} table.",
                    answer=f"SELECT SUM({col['name']}) AS total_{col['name']} FROM {table};",
                    explanation=f"SUM({col['name']}) calculates the total of all {col['name']} values.",
                    tip="SUM ignores NULL values."
                )
                
                self.add_question(
                    topic='Aggregation',
                    difficulty='Intermediate',
                    question=f"Find the maximum {col['name']} from the {table} table.",
                    answer=f"SELECT MAX({col['name']}) AS max_{col['name']} FROM {table};",
                    explanation=f"MAX({col['name']}) finds the highest {col['name']} value.",
                    tip="MAX works on numeric, date, and text columns."
                )
    
    # ============================================================
    # GROUP BY QUESTIONS
    # ============================================================
    
    def generate_group_by_questions(self):
        for table, info in self.tables.items():
            if info['row_count'] < 3:
                continue
            
            group_cols = [col for col in info['columns'] 
                         if col['data_type'] in ['character varying', 'text', 'integer'] 
                         and not col['primary_key']]
            
            for col in group_cols[:2]:
                values = self.db.get_column_values(table, col['name'])
                if len(values) > 1:
                    self.add_question(
                        topic='GROUP BY',
                        difficulty='Intermediate',
                        question=f"Count records in {table} grouped by {col['name']}.",
                        answer=f"SELECT {col['name']}, COUNT(*) FROM {table} GROUP BY {col['name']};",
                        explanation=f"GROUP BY {col['name']} creates groups and COUNT counts each group.",
                        tip="All non-aggregated columns in SELECT must be in GROUP BY."
                    )
                    break
    
    # ============================================================
    # JOIN QUESTIONS
    # ============================================================
    
    def generate_join_questions(self):
        for table, info in self.tables.items():
            for fk in info['foreign_keys']:
                ref_table = fk['ref_table']
                if ref_table in self.tables and self.tables[ref_table]['row_count'] > 0:
                    self.add_question(
                        topic='JOIN',
                        difficulty='Intermediate',
                        question=f"Join the {table} table with the {ref_table} table.",
                        answer=f"SELECT * FROM {table} JOIN {ref_table} ON {table}.{fk['column']} = {ref_table}.{fk['ref_column']};",
                        explanation=f"INNER JOIN connects {table} to {ref_table} on the foreign key relationship.",
                        tip="Use table aliases to make queries more readable."
                    )
    
    # ============================================================
    # SUBQUERY QUESTIONS
    # ============================================================
    
    def generate_subquery_questions(self):
        for table, info in self.tables.items():
            if info['row_count'] > 5:
                num_cols = [col for col in info['columns'] if col['data_type'] in ['integer', 'numeric', 'decimal']]
                if num_cols:
                    col = num_cols[0]
                    self.add_question(
                        topic='Subquery',
                        difficulty='Advanced',
                        question=f"Find records in {table} where {col['name']} is above average.",
                        answer=f"SELECT * FROM {table} WHERE {col['name']} > (SELECT AVG({col['name']}) FROM {table});",
                        explanation=f"Subquery calculates average {col['name']}, main query filters above it.",
                        tip="Subqueries can be slow; consider CTEs for complex queries."
                    )
    
    # ============================================================
    # COMPLEX QUESTIONS
    # ============================================================
    
    def generate_complex_questions(self):
        # Find tables with foreign keys for complex joins
        tables_with_fk = [t for t, info in self.tables.items() if info['foreign_keys'] and info['row_count'] > 0]
        
        if len(tables_with_fk) >= 2:
            table1 = tables_with_fk[0]
            table2 = tables_with_fk[1]
            
            self.add_question(
                topic='Complex Query',
                difficulty='Advanced',
                question=f"Write a query that joins {table1} and {table2} with an aggregate function.",
                answer=f"SELECT {table1}.id, COUNT({table2}.id) FROM {table1} LEFT JOIN {table2} ON {table1}.id = {table2}.{table1}_id GROUP BY {table1}.id;",
                explanation=f"LEFT JOIN ensures all {table1} records appear even without matches.",
                tip="Use LEFT JOIN when you need all rows from the left table."
            )

# ============================================================
# MAIN FUNCTION
# ============================================================

def generate_workbook_from_database(question_count=100):
    print("=" * 60)
    print("📚 Generate Workbook from Database")
    print("=" * 60)
    print()
    
    # Connect to database
    db = DatabaseConnector()
    if not db.connect():
        print("❌ Could not connect to PostgreSQL!")
        return
    
    print("✅ Connected to PostgreSQL")
    
    # Generate questions
    print(f"\n🧠 Generating {question_count} questions from database data...")
    generator = DatabaseQuestionGenerator(db, question_count)
    questions = generator.generate()
    
    print(f"✅ Generated {len(questions)} questions")
    
    # Display sample
    print("\n📝 Sample Questions:")
    for q in questions[:5]:
        print(f"   Q{q['id']}: {q['question']}")
        print(f"      Answer: {q['answer']}")
        print()
    
    # Save to JSON
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    output_file = f"workbook_from_db_{timestamp}.json"
    with open(output_file, 'w') as f:
        json.dump(questions, f, indent=2)
    
    print(f"\n💾 Questions saved to: {output_file}")
    
    # Save to DOCX (optional)
    try:
        from docx import Document
        from docx.shared import Pt
        
        doc = Document()
        doc.add_heading('SQL Workbook from Database', level=1)
        doc.add_paragraph(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        doc.add_paragraph(f'Total Questions: {len(questions)}')
        
        for q in questions:
            doc.add_heading(f"Question {q['id']}", level=2)
            doc.add_paragraph(f"Topic: {q['topic']}")
            doc.add_paragraph(f"Difficulty: {q['difficulty']}")
            doc.add_paragraph(f"Question: {q['question']}")
            doc.add_paragraph(f"Answer: {q['answer']}")
            doc.add_paragraph(f"Explanation: {q['explanation']}")
            doc.add_paragraph(f"Tip: {q['tip']}")
            doc.add_paragraph("─" * 50)
        
        docx_file = f"workbook_from_db_{timestamp}.docx"
        doc.save(docx_file)
        print(f"📝 DOCX workbook saved to: {docx_file}")
    except ImportError:
        print("⚠️  python-docx not installed. Install with: pip install python-docx")
    
    db.disconnect()
    print("\n✅ Done!")

# ============================================================
# RUN
# ============================================================

if __name__ == "__main__":
    import sys
    count = int(sys.argv[1]) if len(sys.argv) > 1 else 100
    generate_workbook_from_database(count)